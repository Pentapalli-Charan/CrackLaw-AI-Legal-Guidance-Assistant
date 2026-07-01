import os
import csv
import json
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List
import xml.etree.ElementTree as ET

logger = logging.getLogger("CrackLaw.Parsers")

class BaseParser(ABC):
    """Abstract Base Class for all document parsers."""
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parses a document and returns the extracted text and structural metadata.
        
        Returns:
            Dict containing:
                "text": Extracted raw text content of the document.
                "metadata": Dict of extracted metadata attributes (title, act_name, section, etc.).
        """
        pass

    def _extract_basic_metadata(self, file_path: str) -> Dict[str, Any]:
        """Helper to extract general file metadata properties."""
        filename = os.path.basename(file_path)
        return {
            "title": os.path.splitext(filename)[0],
            "original_filename": filename,
            "document_type": os.path.splitext(filename)[1].lower().lstrip("."),
            "language": "en",  # Default, can be overridden by specific parsers
            "source": "file",
            "act_name": None,
            "chapter": None,
            "section": None,
            "subsection": None,
            "publication_date": None
        }

class TXTParser(BaseParser):
    """Parser for plain text files."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing TXT file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        return {"text": text, "metadata": metadata}

class PDFParser(BaseParser):
    """Parser for PDF files using pypdf."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing PDF file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required to parse PDF files. Run 'pip install pypdf'.")
            
        text_pages = []
        try:
            reader = pypdf.PdfReader(file_path)
            metadata["page_count"] = len(reader.pages)
            
            # Extract document title if available in PDF metadata
            if reader.metadata and reader.metadata.title:
                metadata["title"] = reader.metadata.title
                
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
                else:
                    # Log page that might need OCR (scanned image)
                    logger.debug("Page %d of PDF %s is empty. OCR may be needed.", page_num + 1, file_path)
                    
            text = "\n--- PAGE BREAK ---\n".join(text_pages)
            
            if not text.strip():
                logger.warning("No text extracted from PDF %s. The file might be scanned/image-only.", file_path)
                # Try simple pytesseract OCR fallback if installed
                text = self._ocr_fallback(file_path)
                
            return {"text": text, "metadata": metadata}
        except Exception as e:
            logger.error("Failed to parse PDF %s: %s", file_path, str(e))
            raise e

    def _ocr_fallback(self, file_path: str) -> str:
        """Fallback method using pytesseract OCR if available on the system."""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            logger.info("Attempting OCR fallback on %s", file_path)
            
            images = convert_from_path(file_path)
            text_pages = []
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                text_pages.append(page_text)
                
            return "\n--- PAGE BREAK ---\n".join(text_pages)
        except Exception as e:
            logger.debug("OCR fallback unavailable or failed: %s", str(e))
            return ""

class DocxParser(BaseParser):
    """Parser for DOCX files using python-docx."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing DOCX file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required to parse DOCX files. Run 'pip install python-docx'.")
            
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            
            # Extract document core properties title if available
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.created:
                metadata["publication_date"] = core_props.created.strftime("%Y-%m-%d")

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
                        
            text = "\n\n".join(paragraphs)
            return {"text": text, "metadata": metadata}
        except Exception as e:
            logger.error("Failed to parse DOCX %s: %s", file_path, str(e))
            raise e

class JSONParser(BaseParser):
    """Parser for JSON files. Formats structured legal files intelligently."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing JSON file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        text = ""
        # If it's a legal QA dataset or standard structure
        if isinstance(data, list):
            formatted_items = []
            for item in data:
                if isinstance(item, dict):
                    # Look for common Q&A fields
                    q = item.get("question") or item.get("prompt") or item.get("query")
                    a = item.get("answer") or item.get("response") or item.get("completion")
                    c = item.get("context") or item.get("text")
                    
                    if q and a:
                        qa_str = f"Question: {q}\nAnswer: {a}"
                        if c:
                            qa_str = f"Context: {c}\n" + qa_str
                        formatted_items.append(qa_str)
                    elif c:
                        formatted_items.append(str(c))
                    else:
                        formatted_items.append(json.dumps(item, indent=2))
                else:
                    formatted_items.append(str(item))
            text = "\n\n---\n\n".join(formatted_items)
            
        elif isinstance(data, dict):
            # Check if dict has content/text field
            text_content = data.get("text") or data.get("content") or data.get("body")
            if text_content:
                text = str(text_content)
                # Propagate metadata from JSON if exists
                for k in ["title", "act_name", "chapter", "section", "subsection", "language", "publication_date"]:
                    if k in data:
                        metadata[k] = data[k]
            else:
                text = json.dumps(data, indent=2)
                
        else:
            text = str(data)
            
        return {"text": text, "metadata": metadata}

class CSVParser(BaseParser):
    """Parser for CSV files."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing CSV file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            
            if headers:
                rows.append("Columns: " + ", ".join(headers))
                for row_idx, row in enumerate(reader):
                    row_parts = []
                    for h, v in zip(headers, row):
                        if v.strip():
                            row_parts.append(f"{h}: {v.strip()}")
                    if row_parts:
                        rows.append(f"Row {row_idx + 1}: " + " | ".join(row_parts))
            else:
                for row_idx, row in enumerate(reader):
                    rows.append(f"Row {row_idx + 1}: " + ", ".join(row))
                    
        text = "\n".join(rows)
        return {"text": text, "metadata": metadata}

class HTMLParser(BaseParser):
    """Parser for HTML files using beautifulsoup4."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing HTML file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required to parse HTML files. Run 'pip install beautifulsoup4'.")
            
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            
        # Extract title
        if soup.title and soup.title.string:
            metadata["title"] = soup.title.string.strip()
            
        # Extract metadata tag values if available
        meta_desc = soup.find("meta", attrs={"name": "description"})
        if meta_desc:
            metadata["description"] = meta_desc.get("content", "")
            
        # Strip script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        text = soup.get_text()
        
        # Clean lines
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        return {"text": text, "metadata": metadata}

class XMLParser(BaseParser):
    """Parser for XML files using xml.etree.ElementTree."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.debug("Parsing XML file: %s", file_path)
        metadata = self._extract_basic_metadata(file_path)
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Simple recursive text extraction
            texts = []
            def recurse_elements(element, depth=0):
                if element.tag:
                    indent = "  " * depth
                    # If it has tag attributes
                    attrs = ""
                    if element.attrib:
                        attrs = " " + " ".join(f"{k}='{v}'" for k, v in element.attrib.items())
                    
                    if element.text and element.text.strip():
                        texts.append(f"{indent}<{element.tag}{attrs}>{element.text.strip()}</{element.tag}>")
                    elif len(element) > 0:
                        texts.append(f"{indent}<{element.tag}{attrs}>")
                        for child in element:
                            recurse_elements(child, depth + 1)
                        texts.append(f"{indent}</{element.tag}>")
                    else:
                        texts.append(f"{indent}<{element.tag}{attrs}/>")
                        
            recurse_elements(root)
            text = "\n".join(texts)
            return {"text": text, "metadata": metadata}
            
        except Exception as e:
            logger.error("Failed to parse XML %s: %s", file_path, str(e))
            raise e

class ParserFactory:
    """Factory class to instantiate the appropriate document parser."""
    
    _parsers: Dict[str, BaseParser] = {
        "txt": TXTParser(),
        "pdf": PDFParser(),
        "docx": DocxParser(),
        "json": JSONParser(),
        "csv": CSVParser(),
        "html": HTMLParser(),
        "htm": HTMLParser(),
        "xml": XMLParser()
    }

    @classmethod
    def get_parser(cls, file_path: str) -> BaseParser:
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        if ext not in cls._parsers:
            raise ValueError(f"No parser implemented for format: '{ext}'. Supported formats: {list(cls._parsers.keys())}")
        return cls._parsers[ext]
